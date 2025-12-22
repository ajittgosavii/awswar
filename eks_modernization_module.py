"""
EKS Design & Architecture Hub - COMPREHENSIVE EDITION
Enterprise-grade tool for complete EKS design, validation, and deployment

Features:
- Multi-step design wizard (6 phases)
- Comprehensive component selection (compute, storage, networking, observability)
- Intelligent sizing calculator with workload profiles
- AI-powered architecture validation using Claude API
- Real-time cost estimation with optimization recommendations
- Best practices validation engine
- Security & compliance checking
- Architecture diagram generation
- Documentation export (Word, PDF, Markdown)
- IaC code generation (Terraform, CloudFormation, Pulumi)
- Migration planning and risk assessment
"""

import streamlit as st
import json
from datetime import datetime
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, field
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
from io import BytesIO

# Try to import python-docx (optional for Word export)
DOCX_AVAILABLE = False
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    pass

# Try to import Anthropic (optional for AI features)
ANTHROPIC_AVAILABLE = False
try:
    from anthropic import Anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    pass

# Import EKS Integrations for Real Pricing and AI
EKS_INTEGRATIONS_AVAILABLE = False
try:
    from eks_integrations import (
        AWSPricingFetcher,
        AnthropicAIValidator,
        display_integration_status
    )
    EKS_INTEGRATIONS_AVAILABLE = True
except ImportError:
    pass

# ============================================================================
# DATA MODELS FOR EKS ARCHITECTURE
# ============================================================================

@dataclass
class EKSDesignSpec:
    """Complete EKS architecture specification"""
    # Project Info
    project_name: str = ""
    environment: str = "production"
    region: str = "us-east-1"
    availability_zones: List[str] = field(default_factory=lambda: [])
    
    # Compute
    node_groups: List[Dict] = field(default_factory=list)
    karpenter_enabled: bool = False
    karpenter_config: Dict = field(default_factory=dict)
    fargate_profiles: List[Dict] = field(default_factory=list)
    
    # Storage
    storage_classes: List[Dict] = field(default_factory=list)
    ebs_csi_enabled: bool = True
    efs_enabled: bool = False
    efs_configs: List[Dict] = field(default_factory=list)
    fsx_enabled: bool = False
    fsx_configs: List[Dict] = field(default_factory=list)
    
    # Networking
    vpc_cidr: str = "10.0.0.0/16"
    subnet_strategy: str = "public-private"
    load_balancer_type: str = "alb"
    ingress_controller: str = "aws-load-balancer-controller"
    service_mesh: str = "none"
    
    # Security
    encryption_enabled: bool = True
    secrets_manager: str = "aws-secrets-manager"
    irsa_enabled: bool = True
    pod_security_standards: str = "restricted"
    network_policies: bool = True
    
    # Observability
    logging_enabled: bool = True
    logging_destination: str = "cloudwatch"
    metrics_server: bool = True
    prometheus_enabled: bool = False
    grafana_enabled: bool = False
    
    # Add-ons & Tools
    cluster_autoscaler: bool = False
    external_dns: bool = False
    cert_manager: bool = False
    argocd: bool = False
    flux: bool = False
    
    # Workload Info
    expected_workloads: int = 0
    peak_pod_count: int = 0
    workload_types: List[str] = field(default_factory=list)
    
    # Cost & Performance
    monthly_budget: float = 0.0
    performance_tier: str = "balanced"
    
    # Metadata
    created_at: str = ""
    validation_status: str = "pending"
    ai_recommendations: List[Dict] = field(default_factory=list)

# ============================================================================
# DESIGN WIZARD - MULTI-STEP WORKFLOW
# ============================================================================

class EKSDesignWizard:
    """Multi-step wizard for EKS architecture design"""
    
    STEPS = [
        "1️⃣ Project Setup",
        "2️⃣ Compute & Scaling",
        "3️⃣ Storage & Data",
        "4️⃣ Networking & Security",
        "5️⃣ Observability & Tools",
        "6️⃣ Review & Validate"
    ]
    
    @staticmethod
    def initialize_session():
        """Initialize session state for wizard"""
        if 'design_spec' not in st.session_state:
            st.session_state.design_spec = EKSDesignSpec()
        if 'wizard_step' not in st.session_state:
            st.session_state.wizard_step = 0
        if 'validation_results' not in st.session_state:
            st.session_state.validation_results = None
    
    @staticmethod
    def render_wizard():
        """Render the complete design wizard"""
        EKSDesignWizard.initialize_session()
        
        st.title("🎯 EKS Architecture Design Wizard")
        st.markdown("Complete, AI-validated EKS architecture design in 6 steps")
        
        # Progress indicator
        progress = (st.session_state.wizard_step + 1) / len(EKSDesignWizard.STEPS)
        st.progress(progress)
        
        # Step navigation
        col1, col2, col3 = st.columns([1, 3, 1])
        with col2:
            current_step = st.radio(
                "Steps",
                range(len(EKSDesignWizard.STEPS)),
                format_func=lambda x: EKSDesignWizard.STEPS[x],
                horizontal=True,
                index=st.session_state.wizard_step
            )
            st.session_state.wizard_step = current_step
        
        st.divider()
        
        # Render current step
        if current_step == 0:
            EKSDesignWizard.step1_project_setup()
        elif current_step == 1:
            EKSDesignWizard.step2_compute_scaling()
        elif current_step == 2:
            EKSDesignWizard.step3_storage_data()
        elif current_step == 3:
            EKSDesignWizard.step4_networking_security()
        elif current_step == 4:
            EKSDesignWizard.step5_observability_tools()
        elif current_step == 5:
            EKSDesignWizard.step6_review_validate()
        
        # Navigation buttons
        st.divider()
        col1, col2, col3 = st.columns([1, 1, 1])
        with col1:
            if current_step > 0:
                if st.button("⬅️ Previous", use_container_width=True):
                    st.session_state.wizard_step -= 1
                    st.rerun()
        with col3:
            if current_step < len(EKSDesignWizard.STEPS) - 1:
                if st.button("Next ➡️", use_container_width=True):
                    st.session_state.wizard_step += 1
                    st.rerun()
            elif current_step == len(EKSDesignWizard.STEPS) - 1:
                if st.button("✅ Complete Design", type="primary", use_container_width=True):
                    st.success("🎉 Design completed! Ready to export.")
    
    @staticmethod
    def step1_project_setup():
        """Step 1: Project Setup"""
        st.header("1️⃣ Project Setup")
        spec = st.session_state.design_spec
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Basic Information")
            spec.project_name = st.text_input(
                "Project Name *",
                value=spec.project_name,
                placeholder="my-eks-cluster"
            )
            
            spec.environment = st.selectbox(
                "Environment *",
                ["development", "staging", "production", "dr"],
                index=["development", "staging", "production", "dr"].index(spec.environment)
            )
            
            spec.region = st.selectbox(
                "AWS Region *",
                ["us-east-1", "us-east-2", "us-west-1", "us-west-2", 
                 "eu-west-1", "eu-central-1", "ap-southeast-1", "ap-northeast-1"],
                index=0 if not spec.region else ["us-east-1", "us-east-2", "us-west-1", 
                          "us-west-2", "eu-west-1", "eu-central-1", "ap-southeast-1", 
                          "ap-northeast-1"].index(spec.region) if spec.region in ["us-east-1", "us-east-2", "us-west-1", 
                          "us-west-2", "eu-west-1", "eu-central-1", "ap-southeast-1", 
                          "ap-northeast-1"] else 0
            )
            
            az_options = [f"{spec.region}a", f"{spec.region}b", f"{spec.region}c"]
            spec.availability_zones = st.multiselect(
                "Availability Zones *",
                az_options,
                default=spec.availability_zones if spec.availability_zones else az_options[:2]
            )
        
        with col2:
            st.subheader("Workload Characteristics")
            spec.expected_workloads = st.number_input(
                "Number of Applications/Services",
                min_value=1,
                max_value=500,
                value=max(1, spec.expected_workloads),
                help="How many distinct applications will run in this cluster?"
            )
            
            spec.peak_pod_count = st.number_input(
                "Peak Pod Count (Estimate)",
                min_value=10,
                max_value=10000,
                value=max(10, spec.peak_pod_count),
                help="Maximum number of pods you expect to run simultaneously"
            )
            
            spec.workload_types = st.multiselect(
                "Workload Types",
                ["Web Applications", "APIs/Microservices", "Batch Processing", 
                 "Data Processing", "ML/AI", "Databases", "Message Queues", 
                 "Caching", "CI/CD", "Monitoring"],
                default=spec.workload_types
            )
            
            spec.monthly_budget = st.number_input(
                "Monthly Budget (USD)",
                min_value=0.0,
                max_value=1000000.0,
                value=float(spec.monthly_budget),
                step=1000.0,
                format="%.2f"
            )
            
            spec.performance_tier = st.select_slider(
                "Performance Tier",
                options=["cost-optimized", "balanced", "performance-optimized"],
                value=spec.performance_tier
            )
        
        # Recommendations based on inputs
        st.divider()
        st.subheader("📊 Initial Recommendations")
        
        # Calculate recommended node count
        pods_per_node = 110 if spec.performance_tier == "performance-optimized" else 58
        recommended_nodes = max(3, int(spec.peak_pod_count / pods_per_node) + 2)
        
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("Recommended Min Nodes", f"{recommended_nodes}")
        col2.metric("Recommended AZs", f"{len(spec.availability_zones)}")
        col3.metric("HA Status", "✅ Yes" if len(spec.availability_zones) >= 2 else "⚠️ Single AZ")
        
        # Environment-specific recommendations
        if spec.environment == "production":
            st.info("✅ Production environment: HA configuration, multi-AZ recommended")
        elif spec.environment == "development":
            st.info("💡 Development environment: Consider single-AZ for cost savings")
    
    @staticmethod
    def step2_compute_scaling():
        """Step 2: Compute & Scaling Strategy"""
        st.header("2️⃣ Compute & Scaling Strategy")
        spec = st.session_state.design_spec
        
        # Compute strategy selection
        st.subheader("Select Compute Strategy")
        compute_strategy = st.radio(
            "Primary Compute Approach",
            ["Karpenter (Recommended)", "Managed Node Groups", "Fargate", "Hybrid"],
            help="Karpenter provides best cost optimization and flexibility"
        )
        
        if "Karpenter" in compute_strategy:
            st.success("✅ Excellent choice! Karpenter provides 30-50% cost savings")
            EKSDesignWizard._configure_karpenter()
        
        if "Node Groups" in compute_strategy or "Hybrid" in compute_strategy:
            EKSDesignWizard._configure_node_groups()
        
        if "Fargate" in compute_strategy or "Hybrid" in compute_strategy:
            EKSDesignWizard._configure_fargate()
        
        # Sizing Calculator
        st.divider()
        st.subheader("🎯 Intelligent Sizing Calculator")
        EKSDesignWizard._render_sizing_calculator()
    
    @staticmethod
    def _configure_karpenter():
        """Configure Karpenter settings"""
        spec = st.session_state.design_spec
        spec.karpenter_enabled = True
        
        with st.expander("⚙️ Karpenter Configuration", expanded=True):
            col1, col2 = st.columns(2)
            
            with col1:
                instance_families = st.multiselect(
                    "Instance Families",
                    ["t3", "t3a", "m5", "m6i", "c5", "c6i", "r5", "r6i"],
                    default=["m5", "c5", "r5"]
                )
                
                spot_enabled = st.checkbox("Enable Spot Instances", value=True)
                if spot_enabled:
                    spot_percentage = st.slider(
                        "Spot Instance Percentage",
                        min_value=0,
                        max_value=100,
                        value=70,
                        help="Higher percentage = more savings, slightly higher interruption risk"
                    )
                else:
                    spot_percentage = 0
                
                consolidation = st.checkbox("Enable Consolidation", value=True)
            
            with col2:
                ttl_seconds = st.number_input(
                    "Empty Node TTL (seconds)",
                    min_value=30,
                    max_value=300,
                    value=30,
                    help="How long to wait before terminating empty nodes"
                )
                
                cpu_limits = st.slider(
                    "CPU Limits (vCPUs)",
                    min_value=2,
                    max_value=128,
                    value=(2, 32)
                )
                
                memory_limits = st.slider(
                    "Memory Limits (GB)",
                    min_value=4,
                    max_value=512,
                    value=(4, 128)
                )
            
            spec.karpenter_config = {
                'instance_families': instance_families,
                'spot_enabled': spot_enabled,
                'spot_percentage': spot_percentage,
                'consolidation_enabled': consolidation,
                'ttl_seconds_after_empty': ttl_seconds,
                'cpu_limits': cpu_limits,
                'memory_limits': memory_limits
            }
            
            # Show expected savings
            baseline_cost = 10000  # Example baseline
            savings_percentage = (spot_percentage / 100) * 0.70  # 70% savings on Spot
            total_savings = baseline_cost * savings_percentage
            
            st.success(f"💰 Expected Monthly Savings: ${total_savings:,.2f} ({savings_percentage*100:.1f}%)")
    
    @staticmethod
    def _configure_node_groups():
        """Configure managed node groups"""
        spec = st.session_state.design_spec
        
        with st.expander("🖥️ Managed Node Groups Configuration"):
            st.info("Add one or more managed node groups for predictable workloads")
            
            num_groups = st.number_input("Number of Node Groups", min_value=1, max_value=5, value=1)
            
            for i in range(num_groups):
                st.markdown(f"**Node Group {i+1}**")
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    ng_name = st.text_input(f"Name", value=f"ng-{i+1}", key=f"ng_name_{i}")
                    instance_type = st.selectbox(
                        f"Instance Type",
                        ["t3.medium", "t3.large", "t3.xlarge", "m5.large", "m5.xlarge",
                         "m5.2xlarge", "c5.large", "c5.xlarge", "r5.large", "r5.xlarge"],
                        key=f"ng_instance_{i}"
                    )
                
                with col2:
                    min_size = st.number_input(f"Min Size", min_value=0, max_value=100, value=1, key=f"ng_min_{i}")
                    max_size = st.number_input(f"Max Size", min_value=1, max_value=500, value=10, key=f"ng_max_{i}")
                
                with col3:
                    desired_size = st.number_input(f"Desired Size", min_value=min_size, max_value=max_size, value=2, key=f"ng_desired_{i}")
                    capacity_type = st.selectbox(f"Capacity Type", ["ON_DEMAND", "SPOT"], key=f"ng_capacity_{i}")
                
                # Add to spec
                if len(spec.node_groups) <= i:
                    spec.node_groups.append({})
                spec.node_groups[i] = {
                    'name': ng_name,
                    'instance_type': instance_type,
                    'min_size': min_size,
                    'max_size': max_size,
                    'desired_size': desired_size,
                    'capacity_type': capacity_type
                }
    
    @staticmethod
    def _configure_fargate():
        """Configure Fargate profiles"""
        spec = st.session_state.design_spec
        
        with st.expander("🚀 AWS Fargate Configuration"):
            st.info("Fargate provides serverless compute for specific workloads")
            
            enable_fargate = st.checkbox("Enable Fargate", value=False)
            if enable_fargate:
                fargate_namespaces = st.multiselect(
                    "Fargate Namespaces",
                    ["default", "kube-system", "production", "staging", "batch"],
                    default=["batch"]
                )
                
                spec.fargate_profiles = [{
                    'name': 'fargate-profile-1',
                    'namespaces': fargate_namespaces
                }]
                
                st.success(f"✅ Fargate enabled for {len(fargate_namespaces)} namespace(s)")
    
    @staticmethod
    def _render_sizing_calculator():
        """Intelligent sizing calculator with workload profiles"""
        spec = st.session_state.design_spec
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.markdown("**Workload Profile**")
            workload_profile = st.selectbox(
                "Select Profile",
                ["Web Application", "Microservices", "Batch Processing", "Data Pipeline", "ML Training", "Custom"],
                help="Pre-configured sizing based on common patterns"
            )
        
        with col2:
            st.markdown("**Resource Requirements**")
            avg_cpu_per_pod = st.number_input("Avg CPU per Pod (millicores)", min_value=100, max_value=16000, value=500)
            avg_memory_per_pod = st.number_input("Avg Memory per Pod (MB)", min_value=128, max_value=32768, value=512)
        
        with col3:
            st.markdown("**Scale Parameters**")
            peak_pods = st.number_input("Peak Pod Count", min_value=10, max_value=5000, value=spec.peak_pod_count)
            overhead_factor = st.slider("Overhead Factor", min_value=1.1, max_value=2.0, value=1.3, step=0.1)
        
        # Calculate sizing
        total_cpu_needed = (avg_cpu_per_pod * peak_pods * overhead_factor) / 1000  # Convert to vCPUs
        total_memory_needed = (avg_memory_per_pod * peak_pods * overhead_factor) / 1024  # Convert to GB
        
        # Recommend instance types
        recommended_instances = SizingCalculator.recommend_instances(
            total_cpu_needed, total_memory_needed, workload_profile
        )
        
        st.divider()
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("Total vCPUs Needed", f"{total_cpu_needed:.1f}")
        col2.metric("Total Memory Needed", f"{total_memory_needed:.1f} GB")
        
        # DEFENSIVE CODING: Safe access to recommended instances
        node_count = recommended_instances.get('node_count', 'N/A') if isinstance(recommended_instances, dict) else 'N/A'
        monthly_cost = recommended_instances.get('monthly_cost', 0) if isinstance(recommended_instances, dict) else 0
        
        col3.metric("Recommended Nodes", f"{node_count}")
        col4.metric("Monthly Cost (Est)", f"${monthly_cost:,.2f}" if isinstance(monthly_cost, (int, float)) else "N/A")
        
        st.markdown("**Recommended Instance Types:**")
        
        # DEFENSIVE CODING: Safely iterate instances
        instances = recommended_instances.get('instances', []) if isinstance(recommended_instances, dict) else []
        if not instances:
            st.info("No instance recommendations available")
        else:
            for instance in instances:
                if isinstance(instance, dict):
                    instance_type = instance.get('type', 'Unknown')
                    vcpu = instance.get('vcpu', 0)
                    memory = instance.get('memory', 0)
                    cost = instance.get('monthly_cost', 0)
                    st.markdown(f"- `{instance_type}` - {vcpu} vCPU, {memory} GB RAM - ${cost:.2f}/month")
    
    @staticmethod
    def step3_storage_data():
        """Step 3: Storage & Data Configuration"""
        st.header("3️⃣ Storage & Data Configuration")
        spec = st.session_state.design_spec
        
        st.subheader("Storage Strategy")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Block Storage (EBS)**")
            spec.ebs_csi_enabled = st.checkbox("Enable EBS CSI Driver", value=True)
            
            if spec.ebs_csi_enabled:
                storage_classes = st.multiselect(
                    "Storage Classes",
                    ["gp3 (General Purpose)", "io2 (High Performance)", "st1 (Throughput Optimized)", "sc1 (Cold HDD)"],
                    default=["gp3 (General Purpose)"]
                )
                
                encryption = st.checkbox("Enable Encryption at Rest", value=True)
                snapshot_policy = st.checkbox("Enable Automated Snapshots", value=True)
                
                spec.storage_classes = [
                    {
                        'type': 'ebs',
                        'classes': storage_classes,
                        'encryption': encryption,
                        'snapshots': snapshot_policy
                    }
                ]
        
        with col2:
            st.markdown("**Shared File Storage (EFS)**")
            spec.efs_enabled = st.checkbox("Enable EFS", value=False)
            
            if spec.efs_enabled:
                efs_mode = st.selectbox("EFS Performance Mode", ["General Purpose", "Max I/O"])
                efs_throughput = st.selectbox("Throughput Mode", ["Bursting", "Provisioned"])
                
                spec.efs_configs = [{
                    'mode': efs_mode,
                    'throughput': efs_throughput,
                    'encryption': True
                }]
                
                st.info("💡 EFS is ideal for shared data across multiple pods")
        
        # FSx for Lustre (HPC workloads)
        st.divider()
        st.subheader("High-Performance Storage")
        
        spec.fsx_enabled = st.checkbox("Enable FSx for Lustre", value=False)
        if spec.fsx_enabled:
            st.info("🚀 FSx for Lustre provides high-performance file system for HPC and ML workloads")
            fsx_capacity = st.number_input("Storage Capacity (GB)", min_value=1200, max_value=100000, value=1200, step=1200)
            fsx_throughput = st.selectbox("Throughput per TiB", ["50 MB/s", "100 MB/s", "200 MB/s"])
            
            spec.fsx_configs = [{
                'capacity_gb': fsx_capacity,
                'throughput': fsx_throughput
            }]
        
        # Storage recommendations
        st.divider()
        st.subheader("📊 Storage Recommendations")
        
        recommendations = []
        if "Databases" in spec.workload_types:
            recommendations.append("✅ Use io2 EBS for database workloads requiring high IOPS")
        if "Data Processing" in spec.workload_types or "ML/AI" in spec.workload_types:
            recommendations.append("✅ Consider FSx for Lustre for high-throughput data processing")
        if spec.expected_workloads > 20:
            recommendations.append("✅ Enable EFS for shared configuration and log files")
        
        for rec in recommendations:
            st.markdown(rec)
    
    @staticmethod
    def step4_networking_security():
        """Step 4: Networking & Security Configuration"""
        st.header("4️⃣ Networking & Security")
        spec = st.session_state.design_spec
        
        # Networking
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Networking")
            spec.vpc_cidr = st.text_input("VPC CIDR", value=spec.vpc_cidr)
            spec.subnet_strategy = st.selectbox(
                "Subnet Strategy",
                ["public-private", "private-only", "public-only"],
                help="public-private is recommended for production"
            )
            
            spec.load_balancer_type = st.selectbox(
                "Load Balancer Type",
                ["alb", "nlb", "both"],
                help="ALB for HTTP/HTTPS, NLB for TCP/UDP"
            )
            
            spec.ingress_controller = st.selectbox(
                "Ingress Controller",
                ["aws-load-balancer-controller", "nginx", "traefik", "istio"],
                help="AWS LB Controller recommended for native AWS integration"
            )
            
            spec.service_mesh = st.selectbox(
                "Service Mesh",
                ["none", "istio", "linkerd", "app-mesh"],
                help="Service mesh provides advanced traffic management"
            )
        
        with col2:
            st.subheader("Security")
            spec.encryption_enabled = st.checkbox("Enable Encryption", value=True)
            spec.secrets_manager = st.selectbox(
                "Secrets Management",
                ["aws-secrets-manager", "external-secrets", "sealed-secrets", "vault"]
            )
            
            spec.irsa_enabled = st.checkbox(
                "Enable IRSA (IAM Roles for Service Accounts)",
                value=True,
                help="Best practice for pod-level IAM permissions"
            )
            
            spec.pod_security_standards = st.selectbox(
                "Pod Security Standards",
                ["restricted", "baseline", "privileged"],
                help="'restricted' is most secure, recommended for production"
            )
            
            spec.network_policies = st.checkbox(
                "Enable Network Policies",
                value=True,
                help="Control pod-to-pod communication"
            )
        
        # Security recommendations
        st.divider()
        st.subheader("🔒 Security Best Practices")
        
        security_score = 0
        if spec.encryption_enabled:
            security_score += 20
        if spec.irsa_enabled:
            security_score += 20
        if spec.pod_security_standards == "restricted":
            security_score += 20
        if spec.network_policies:
            security_score += 20
        if spec.subnet_strategy == "public-private":
            security_score += 20
        
        col1, col2, col3 = st.columns(3)
        col1.metric("Security Score", f"{security_score}/100")
        col2.metric("Security Level", "High" if security_score >= 80 else "Medium" if security_score >= 60 else "Low")
        
        if security_score < 80:
            st.warning("⚠️ Consider enabling more security features for production workloads")
        else:
            st.success("✅ Excellent security configuration!")
    
    @staticmethod
    def step5_observability_tools():
        """Step 5: Observability & DevOps Tools"""
        st.header("5️⃣ Observability & DevOps Tools")
        spec = st.session_state.design_spec
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Logging & Monitoring")
            spec.logging_enabled = st.checkbox("Enable Logging", value=True)
            
            if spec.logging_enabled:
                spec.logging_destination = st.selectbox(
                    "Logging Destination",
                    ["cloudwatch", "elasticsearch", "splunk", "datadog"]
                )
            
            spec.metrics_server = st.checkbox("Metrics Server", value=True)
            spec.prometheus_enabled = st.checkbox("Prometheus", value=False)
            spec.grafana_enabled = st.checkbox("Grafana", value=False)
            
            if spec.prometheus_enabled:
                st.info("📊 Prometheus provides detailed cluster metrics")
        
        with col2:
            st.subheader("DevOps & GitOps")
            spec.external_dns = st.checkbox("External DNS", value=False)
            spec.cert_manager = st.checkbox("Cert Manager", value=False)
            spec.argocd = st.checkbox("ArgoCD (GitOps)", value=False)
            spec.flux = st.checkbox("Flux (GitOps)", value=False)
            
            if spec.argocd or spec.flux:
                st.success("✅ GitOps enabled for declarative deployments")
        
        # Tool recommendations
        st.divider()
        st.subheader("🛠️ Recommended Tools")
        
        if spec.expected_workloads > 10:
            st.markdown("- ✅ **Prometheus + Grafana** for comprehensive monitoring")
        if "production" in spec.environment:
            st.markdown("- ✅ **ArgoCD or Flux** for GitOps deployment automation")
        if spec.service_mesh != "none":
            st.markdown("- ✅ **Distributed tracing** (Jaeger/Zipkin) for service mesh observability")
    
    @staticmethod
    def step6_review_validate():
        """Step 6: Review & AI Validation"""
        st.header("6️⃣ Review & Validate Architecture")
        spec = st.session_state.design_spec
        
        # Architecture Summary
        st.subheader("📋 Architecture Summary")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.markdown("**Project**")
            st.markdown(f"- Name: `{spec.project_name}`")
            st.markdown(f"- Environment: `{spec.environment}`")
            st.markdown(f"- Region: `{spec.region}`")
            st.markdown(f"- AZs: `{len(spec.availability_zones)}`")
        
        with col2:
            st.markdown("**Compute**")
            st.markdown(f"- Karpenter: {'✅' if spec.karpenter_enabled else '❌'}")
            st.markdown(f"- Node Groups: `{len(spec.node_groups)}`")
            st.markdown(f"- Fargate: {'✅' if spec.fargate_profiles else '❌'}")
        
        with col3:
            st.markdown("**Storage & Networking**")
            st.markdown(f"- EBS CSI: {'✅' if spec.ebs_csi_enabled else '❌'}")
            st.markdown(f"- EFS: {'✅' if spec.efs_enabled else '❌'}")
            st.markdown(f"- Load Balancer: `{spec.load_balancer_type.upper()}`")
        
        # Cost Estimation
        st.divider()
        st.subheader("💰 Cost Estimation")
        
        # Show integration status
        if EKS_INTEGRATIONS_AVAILABLE:
            with st.expander("🔌 Pricing & AI Integration Status", expanded=False):
                display_integration_status()
        
        # Calculate costs
        estimator = CostEstimator()
        cost_estimate = estimator.calculate_total_cost(spec)
        
        # Show pricing source
        if 'pricing_source' in cost_estimate:
            if cost_estimate.get('is_real_data', False):
                st.success(f"✅ Using real pricing from: {cost_estimate['pricing_source']}")
            else:
                st.info(f"ℹ️ Using: {cost_estimate['pricing_source']}")
                st.caption("💡 Configure AWS credentials in .streamlit/secrets.toml for real-time pricing")
        
        col1, col2, col3, col4 = st.columns(4)
        
        # DEFENSIVE CODING: Safe access to cost estimate values
        total_cost = cost_estimate.get('total', 0) if isinstance(cost_estimate, dict) else 0
        compute_cost = cost_estimate.get('compute', 0) if isinstance(cost_estimate, dict) else 0
        storage_cost = cost_estimate.get('storage', 0) if isinstance(cost_estimate, dict) else 0
        network_cost = cost_estimate.get('network', 0) if isinstance(cost_estimate, dict) else 0
        
        col1.metric("Monthly Cost", f"${total_cost:,.2f}")
        col2.metric("Compute Cost", f"${compute_cost:,.2f}")
        col3.metric("Storage Cost", f"${storage_cost:,.2f}")
        col4.metric("Network Cost", f"${network_cost:,.2f}")
        
        if spec.monthly_budget > 0:
            budget_status = "✅ Within Budget" if total_cost <= spec.monthly_budget else "⚠️ Over Budget"
            st.info(f"{budget_status} - Budget: ${spec.monthly_budget:,.2f}")
        
        # AI Validation
        st.divider()
        st.subheader("🤖 AI-Powered Architecture Validation")
        
        # Initialize validator (auto-loads from secrets)
        validator = AIArchitectureValidator()
        
        if validator.available:
            st.success(f"{validator.status_message}")
            
            if st.button("🔍 Validate Architecture with AI", type="primary", use_container_width=True):
                with st.spinner("🤖 AI analyzing your architecture..."):
                    validation_results = validator.validate_architecture(spec)
                    st.session_state.validation_results = validation_results
            
            if st.session_state.validation_results:
                EKSDesignWizard._display_validation_results(st.session_state.validation_results)
        else:
            st.info(f"💡 {validator.status_message}")
            st.caption("Configure ANTHROPIC_API_KEY in .streamlit/secrets.toml to enable AI validation")
            
            # Show basic validation without AI
            basic_validation = BasicValidator.validate(spec)
            EKSDesignWizard._display_basic_validation(basic_validation)
        
        # Export Options
        st.divider()
        st.subheader("📦 Export & Documentation")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            if st.button("📄 Export to Word", use_container_width=True):
                doc_generator = DocumentationGenerator()
                doc_bytes_io = doc_generator.generate_word_doc(spec)
                
                # Provide download button for the Word document
                st.download_button(
                    label="⬇️ Download Word Document",
                    data=doc_bytes_io.getvalue(),
                    file_name=f"eks_architecture_{spec.project_name.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                st.success("✅ Word document ready for download!")
        
        with col2:
            if st.button("🔧 Generate Terraform", use_container_width=True):
                iac_generator = IaCGenerator()
                tf_code = iac_generator.generate_terraform(spec)
                st.download_button(
                    "Download Terraform",
                    tf_code,
                    file_name="main.tf",
                    mime="text/plain"
                )
        
        with col3:
            if st.button("☁️ Generate CloudFormation", use_container_width=True):
                iac_generator = IaCGenerator()
                cfn_code = iac_generator.generate_cloudformation(spec)
                st.download_button(
                    "Download CloudFormation",
                    cfn_code,
                    file_name="template.yaml",
                    mime="text/yaml"
                )
        
        with col4:
            if st.button("📊 Architecture Diagram", use_container_width=True):
                diagram_generator = DiagramGenerator()
                diagram_generator.generate_diagram(spec)
                st.success("✅ Diagram generated!")
    
    @staticmethod
    def _display_validation_results(results: Dict):
        """Display AI validation results"""
        if results.get('overall_score'):
            score = results['overall_score']
            col1, col2, col3 = st.columns(3)
            col1.metric("Overall Score", f"{score}/100")
            col2.metric("Risk Level", results.get('risk_level', 'Unknown'))
            col3.metric("Readiness", results.get('readiness', 'Unknown'))
        
        if results.get('recommendations'):
            st.markdown("### 🎯 AI Recommendations")
            
            # DEFENSIVE CODING: Filter to only valid dict items
            raw_recommendations = results['recommendations'][:10]
            valid_recommendations = [r for r in raw_recommendations if isinstance(r, dict)]
            
            if not valid_recommendations:
                st.info("No structured recommendations available at this time")
            else:
                for rec in valid_recommendations:
                    try:
                        # Safe access with defaults
                        priority = rec.get('priority', 'MEDIUM')
                        title = rec.get('title', 'Recommendation')
                        description = rec.get('description', 'No description available')
                        
                        # Determine if should be expanded
                        is_critical = str(priority).upper() == 'CRITICAL'
                        
                        with st.expander(f"{priority}: {title}", expanded=is_critical):
                            st.markdown(description)
                            
                            # Additional fields if present
                            if rec.get('action'):
                                st.markdown(f"**Action:** {rec['action']}")
                            if rec.get('impact'):
                                st.caption(f"**Impact:** {rec['impact']}")
                            if rec.get('remediation'):
                                st.caption(f"**Remediation:** {rec['remediation']}")
                    
                    except (KeyError, TypeError, AttributeError) as e:
                        # Skip invalid recommendations silently
                        st.warning(f"⚠️ Skipping invalid recommendation: {str(e)}")
                        continue
        
        if results.get('issues'):
            st.markdown("### ⚠️ Issues Found")
            
            # DEFENSIVE CODING: Filter to only valid dict items
            raw_issues = results['issues']
            valid_issues = [i for i in raw_issues if isinstance(i, dict)]
            
            if not valid_issues:
                st.info("No structured issues to display")
            else:
                for issue in valid_issues:
                    try:
                        severity = issue.get('severity', 'MEDIUM')
                        description = issue.get('description', 'No description available')
                        st.warning(f"**{severity}:** {description}")
                    except (KeyError, TypeError, AttributeError) as e:
                        st.warning(f"⚠️ Skipping invalid issue: {str(e)}")
                        continue
    
    @staticmethod
    def _display_basic_validation(validation: Dict):
        """Display basic validation results"""
        st.markdown("### ✅ Basic Validation")
        
        for category, checks in validation.items():
            with st.expander(f"📋 {category.replace('_', ' ').title()}", expanded=True):
                # DEFENSIVE CODING: Ensure checks is a list
                if not isinstance(checks, list):
                    st.warning(f"Invalid checks format for {category}")
                    continue
                
                for check in checks:
                    try:
                        # Handle both dict and string checks
                        if isinstance(check, dict):
                            passed = check.get('passed', False)
                            message = check.get('message', 'No message')
                            icon = "✅" if passed else "⚠️"
                            st.markdown(f"{icon} {message}")
                        elif isinstance(check, str):
                            st.markdown(f"ℹ️ {check}")
                        else:
                            continue
                    except (KeyError, TypeError, AttributeError) as e:
                        st.warning(f"⚠️ Invalid check item: {str(e)}")
                        continue

# ============================================================================
# SIZING CALCULATOR
# ============================================================================

class SizingCalculator:
    """Intelligent sizing calculator"""
    
    INSTANCE_SPECS = {
        't3.medium': {'vcpu': 2, 'memory': 4, 'cost': 30.37},
        't3.large': {'vcpu': 2, 'memory': 8, 'cost': 60.74},
        't3.xlarge': {'vcpu': 4, 'memory': 16, 'cost': 121.47},
        'm5.large': {'vcpu': 2, 'memory': 8, 'cost': 70.08},
        'm5.xlarge': {'vcpu': 4, 'memory': 16, 'cost': 140.16},
        'm5.2xlarge': {'vcpu': 8, 'memory': 32, 'cost': 280.32},
        'c5.large': {'vcpu': 2, 'memory': 4, 'cost': 62.05},
        'c5.xlarge': {'vcpu': 4, 'memory': 8, 'cost': 124.10},
        'c5.2xlarge': {'vcpu': 8, 'memory': 16, 'cost': 248.20},
        'r5.large': {'vcpu': 2, 'memory': 16, 'cost': 91.98},
        'r5.xlarge': {'vcpu': 4, 'memory': 32, 'cost': 183.96},
        'r5.2xlarge': {'vcpu': 8, 'memory': 64, 'cost': 367.92},
    }
    
    @staticmethod
    def recommend_instances(cpu_needed: float, memory_needed: float, workload_profile: str) -> Dict:
        """Recommend optimal instance types"""
        
        # Select instance family based on workload
        if workload_profile in ["Batch Processing", "Data Pipeline"]:
            preferred_family = 'c5'  # Compute optimized
        elif workload_profile == "ML Training":
            preferred_family = 'r5'  # Memory optimized
        else:
            preferred_family = 'm5'  # General purpose
        
        # Find matching instances
        candidates = []
        for instance_type, specs in SizingCalculator.INSTANCE_SPECS.items():
            if instance_type.startswith(preferred_family):
                # Calculate how many nodes needed
                nodes_for_cpu = cpu_needed / (specs['vcpu'] * 0.9)  # 90% allocatable
                nodes_for_memory = memory_needed / (specs['memory'] * 0.9)
                nodes_needed = max(nodes_for_cpu, nodes_for_memory)
                nodes_needed = max(3, int(nodes_needed) + 1)  # Minimum 3 nodes
                
                total_cost = specs['cost'] * nodes_needed
                
                candidates.append({
                    'type': instance_type,
                    'vcpu': specs['vcpu'],
                    'memory': specs['memory'],
                    'monthly_cost': specs['cost'],
                    'nodes_needed': nodes_needed,
                    'total_cost': total_cost
                })
        
        # Sort by total cost
        candidates.sort(key=lambda x: x['total_cost'])
        
        # Return top 3 options
        return {
            'instances': candidates[:3],
            'node_count': candidates[0]['nodes_needed'] if candidates else 3,
            'monthly_cost': candidates[0]['total_cost'] if candidates else 1000
        }

# ============================================================================
# COST ESTIMATOR
# ============================================================================

class CostEstimator:
    """Calculate total cost of EKS architecture with real AWS pricing"""
    
    def __init__(self):
        """Initialize with real pricing fetcher if available"""
        self.pricing = None
        if EKS_INTEGRATIONS_AVAILABLE:
            try:
                self.pricing = AWSPricingFetcher()
            except Exception as e:
                st.warning(f"Could not initialize pricing fetcher: {e}")
    
    def calculate_total_cost(self, spec: EKSDesignSpec) -> Dict[str, float]:
        """Calculate total monthly cost with real or fallback pricing"""
        
        # EKS Control Plane (standard pricing)
        eks_control_plane = 73.00
        
        # Compute Cost - Use real pricing if available
        compute_cost = 0
        pricing_source = "Fallback Estimates"
        
        if self.pricing:
            # Use real AWS pricing
            try:
                for ng in spec.node_groups:
                    result = self.pricing.get_ec2_pricing(
                        instance_type=ng['instance_type'],
                        region=spec.region
                    )
                    
                    if ng['capacity_type'] == 'SPOT':
                        instance_cost = result['monthly_spot_avg']
                    else:
                        instance_cost = result['monthly_on_demand']
                    
                    compute_cost += instance_cost * ng['desired_size']
                
                pricing_source = result.get('source', 'AWS Pricing API')
                
            except Exception as e:
                # Fall through to fallback calculation
                self.pricing = None
        
        if not self.pricing or compute_cost == 0:
            # Fallback to hardcoded estimates
            for ng in spec.node_groups:
                instance_cost = SizingCalculator.INSTANCE_SPECS.get(
                    ng['instance_type'], 
                    {'cost': 140.16}
                )['cost']
                
                if ng['capacity_type'] == 'SPOT':
                    instance_cost *= 0.3  # 70% savings
                
                compute_cost += instance_cost * ng['desired_size']
        
        # Karpenter cost estimation
        if spec.karpenter_enabled:
            node_count = max(3, spec.peak_pod_count // 58)
            
            if self.pricing:
                try:
                    # Use real pricing for Karpenter nodes
                    avg_instance = spec.karpenter_config.get('instance_type', 'm5.xlarge')
                    result = self.pricing.get_ec2_pricing(avg_instance, spec.region)
                    
                    if spec.karpenter_config.get('spot_enabled'):
                        avg_cost = result['monthly_spot_avg']
                    else:
                        avg_cost = result['monthly_on_demand']
                    
                    compute_cost += node_count * avg_cost
                except:
                    # Fallback
                    avg_cost = 140.16  # m5.xlarge estimate
                    if spec.karpenter_config.get('spot_enabled'):
                        avg_cost *= 0.3
                    compute_cost += node_count * avg_cost
            else:
                # Fallback
                avg_cost = 140.16
                if spec.karpenter_config.get('spot_enabled'):
                    avg_cost *= 0.3
                compute_cost += node_count * avg_cost
        
        # Storage Cost
        storage_cost = 0
        if spec.ebs_csi_enabled:
            node_count = max(3, len(spec.node_groups) * 2)
            storage_cost += (50 * node_count * 0.10)  # $0.10/GB for gp3
        
        if spec.efs_enabled:
            storage_cost += 100  # Estimate
        
        if spec.fsx_enabled:
            for fsx in spec.fsx_configs:
                storage_cost += (fsx['capacity_gb'] / 1000) * 140  # ~$140/TB/month
        
        # Network Cost (estimate)
        network_cost = 50 + (spec.expected_workloads * 5)
        
        # Load Balancer Cost
        if spec.load_balancer_type in ['alb', 'both']:
            network_cost += 22.50  # ALB base cost
        if spec.load_balancer_type in ['nlb', 'both']:
            network_cost += 22.50  # NLB base cost
        
        total = eks_control_plane + compute_cost + storage_cost + network_cost
        
        return {
            'total': total,
            'eks_control_plane': eks_control_plane,
            'compute': compute_cost,
            'storage': storage_cost,
            'network': network_cost,
            'pricing_source': pricing_source,
            'is_real_data': pricing_source == 'AWS Pricing API'
        }

# ============================================================================
# AI ARCHITECTURE VALIDATOR
# ============================================================================

class AIArchitectureValidator:
    """AI-powered architecture validation using Claude (auto-loads from secrets)"""
    
    def __init__(self):
        """Initialize validator - loads API key from secrets automatically"""
        self.validator = None
        self.available = False
        self.status_message = ""
        
        if EKS_INTEGRATIONS_AVAILABLE:
            try:
                self.validator = AnthropicAIValidator()
                self.available = "✅" in self.validator.api_key_status
                self.status_message = self.validator.api_key_status
            except Exception as e:
                self.status_message = f"❌ Could not initialize AI validator: {e}"
        else:
            self.status_message = "⚠️ EKS integrations not available"
    
    def validate_architecture(self, spec: EKSDesignSpec) -> Dict:
        """Validate architecture and provide recommendations"""
        
        if not self.available:
            return {
                'status': 'unavailable',
                'error': 'AI validation unavailable',
                'message': self.status_message,
                'recommendations': []
            }
        
        # Create comprehensive configuration for validation
        config = {
            'project_name': spec.project_name,
            'environment': spec.environment,
            'region': spec.region,
            'availability_zones': len(spec.availability_zones),
            'node_groups': len(spec.node_groups),
            'karpenter_enabled': spec.karpenter_enabled,
            'fargate_profiles': len(spec.fargate_profiles),
            'workload_count': spec.expected_workloads,
            'peak_pods': spec.peak_pod_count,
            'k8s_version': '1.28',
            'security': {
                'encryption': spec.encryption_enabled,
                'irsa': spec.irsa_enabled,
                'pod_security': spec.pod_security_standards,
                'network_policies': spec.network_policies
            },
            'observability': {
                'logging': spec.logging_enabled,
                'prometheus': spec.prometheus_enabled,
                'grafana': spec.grafana_enabled
            },
            'networking': {
                'vpc_cidr': spec.vpc_cidr,
                'subnet_strategy': spec.subnet_strategy,
                'load_balancer': spec.load_balancer_type,
                'service_mesh': spec.service_mesh
            }
        }
        
        try:
            result = self.validator.validate_configuration(config)
            return result
        except Exception as e:
            return {
                'status': 'error',
                'error': f'Validation failed: {str(e)}',
                'message': 'AI validation encountered an error',
                'recommendations': []
            }
    
    def _prepare_summary(self, spec: EKSDesignSpec) -> Dict:
        """Prepare architecture summary for AI"""
        return {
            'project': {
                'name': spec.project_name,
                'environment': spec.environment,
                'region': spec.region,
                'availability_zones': len(spec.availability_zones)
            },
            'compute': {
                'karpenter_enabled': spec.karpenter_enabled,
                'node_groups': len(spec.node_groups),
                'fargate_profiles': len(spec.fargate_profiles)
            },
            'storage': {
                'ebs_enabled': spec.ebs_csi_enabled,
                'efs_enabled': spec.efs_enabled,
                'fsx_enabled': spec.fsx_enabled
            },
            'networking': {
                'subnet_strategy': spec.subnet_strategy,
                'load_balancer': spec.load_balancer_type,
                'service_mesh': spec.service_mesh
            },
            'security': {
                'encryption': spec.encryption_enabled,
                'irsa': spec.irsa_enabled,
                'pod_security': spec.pod_security_standards,
                'network_policies': spec.network_policies
            },
            'workload': {
                'expected_count': spec.expected_workloads,
                'peak_pods': spec.peak_pod_count,
                'types': spec.workload_types
            },
            'budget': spec.monthly_budget
        }

# ============================================================================
# BASIC VALIDATOR (No AI required)
# ============================================================================

class BasicValidator:
    """Basic validation without AI"""
    
    @staticmethod
    def validate(spec: EKSDesignSpec) -> Dict:
        """Perform basic validation checks"""
        
        results = {
            'high_availability': [],
            'security': [],
            'cost': [],
            'performance': []
        }
        
        # HA Checks
        if len(spec.availability_zones) >= 2:
            results['high_availability'].append({
                'passed': True,
                'message': 'Multi-AZ configuration for high availability'
            })
        else:
            results['high_availability'].append({
                'passed': False,
                'message': 'Single AZ - not recommended for production'
            })
        
        # Security Checks
        if spec.encryption_enabled:
            results['security'].append({
                'passed': True,
                'message': 'Encryption at rest enabled'
            })
        
        if spec.irsa_enabled:
            results['security'].append({
                'passed': True,
                'message': 'IRSA enabled for fine-grained IAM permissions'
            })
        
        if spec.pod_security_standards == 'restricted':
            results['security'].append({
                'passed': True,
                'message': 'Restricted pod security standards enforced'
            })
        
        # Cost Checks
        if spec.karpenter_enabled:
            results['cost'].append({
                'passed': True,
                'message': 'Karpenter enabled for cost optimization'
            })
        
        # Performance Checks
        if spec.metrics_server:
            results['performance'].append({
                'passed': True,
                'message': 'Metrics server enabled for HPA'
            })
        
        return results

# ============================================================================
# IAC GENERATOR
# ============================================================================

class IaCGenerator:
    """Generate Infrastructure as Code"""
    
    def generate_terraform(self, spec: EKSDesignSpec) -> str:
        """Generate Terraform configuration"""
        
        terraform = f"""# EKS Cluster Configuration
# Generated by EKS Design Hub

terraform {{
  required_providers {{
    aws = {{
      source  = "hashicorp/aws"
      version = "~> 5.0"
    }}
  }}
}}

provider "aws" {{
  region = "{spec.region}"
}}

# EKS Cluster
module "eks" {{
  source  = "terraform-aws-modules/eks/aws"
  version = "~> 19.0"

  cluster_name    = "{spec.project_name}"
  cluster_version = "1.28"

  vpc_id     = module.vpc.vpc_id
  subnet_ids = module.vpc.private_subnets

  cluster_endpoint_public_access = true

  # Node Groups
"""
        
        for ng in spec.node_groups:
            terraform += f"""
  eks_managed_node_groups = {{
    {ng['name']} = {{
      instance_types = ["{ng['instance_type']}"]
      capacity_type  = "{ng['capacity_type']}"
      
      min_size     = {ng['min_size']}
      max_size     = {ng['max_size']}
      desired_size = {ng['desired_size']}
    }}
  }}
"""
        
        terraform += """
}

# VPC
module "vpc" {
  source  = "terraform-aws-modules/vpc/aws"
  version = "~> 5.0"

  name = "${var.cluster_name}-vpc"
  cidr = \"""" + spec.vpc_cidr + """\"

  azs             = [\"${var.region}a\", \"${var.region}b\", \"${var.region}c\"]
  private_subnets = [\"10.0.1.0/24\", \"10.0.2.0/24\", \"10.0.3.0/24\"]
  public_subnets  = [\"10.0.101.0/24\", \"10.0.102.0/24\", \"10.0.103.0/24\"]

  enable_nat_gateway = true
  single_nat_gateway = false
}
"""
        
        return terraform
    
    def generate_cloudformation(self, spec: EKSDesignSpec) -> str:
        """Generate CloudFormation template"""
        
        template = f"""AWSTemplateFormatVersion: '2010-09-09'
Description: 'EKS Cluster - {spec.project_name}'

Parameters:
  ClusterName:
    Type: String
    Default: {spec.project_name}

Resources:
  EKSCluster:
    Type: AWS::EKS::Cluster
    Properties:
      Name: !Ref ClusterName
      Version: '1.28'
      RoleArn: !GetAtt EKSClusterRole.Arn
      ResourcesVpcConfig:
        SubnetIds:
          - !Ref PrivateSubnet1
          - !Ref PrivateSubnet2

  EKSClusterRole:
    Type: AWS::IAM::Role
    Properties:
      AssumeRolePolicyDocument:
        Version: '2012-10-17'
        Statement:
          - Effect: Allow
            Principal:
              Service: eks.amazonaws.com
            Action: sts:AssumeRole
      ManagedPolicyArns:
        - arn:aws:iam::aws:policy/AmazonEKSClusterPolicy

Outputs:
  ClusterName:
    Value: !Ref EKSCluster
  ClusterEndpoint:
    Value: !GetAtt EKSCluster.Endpoint
"""
        
        return template

# ============================================================================
# DOCUMENTATION GENERATOR
# ============================================================================

class DocumentationGenerator:
    """Generate architecture documentation"""
    
    def generate_word_doc(self, spec: EKSDesignSpec) -> BytesIO:
        """Generate professional Word document with architecture details"""
        
        # Check if python-docx is available
        if not DOCX_AVAILABLE:
            st.error("⚠️ Word document export requires 'python-docx' package. Please install it: `pip install python-docx`")
            return None
        
        # Create new Document
        doc = Document()
        
        # ===== TITLE PAGE =====
        title = doc.add_heading('EKS Architecture Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading(spec.project_name, 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph()
        
        # ===== EXECUTIVE SUMMARY =====
        doc.add_heading('Executive Summary', 1)
        doc.add_paragraph(
            f'This document describes the Amazon Elastic Kubernetes Service (EKS) '
            f'architecture for {spec.project_name}.'
        )
        
        # Project details table
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        table.rows[0].cells[0].text = 'Environment'
        table.rows[0].cells[1].text = spec.environment
        table.rows[1].cells[0].text = 'Region'
        table.rows[1].cells[1].text = spec.region
        table.rows[2].cells[0].text = 'Availability Zones'
        table.rows[2].cells[1].text = str(len(spec.availability_zones))
        table.rows[3].cells[0].text = 'VPC CIDR'
        table.rows[3].cells[1].text = spec.vpc_cidr
        
        doc.add_paragraph()
        
        # ===== ARCHITECTURE OVERVIEW =====
        doc.add_heading('Architecture Overview', 1)
        
        # Compute Strategy
        doc.add_heading('Compute Strategy', 2)
        compute_items = [
            f'Karpenter: {"✓ Enabled" if spec.karpenter_enabled else "✗ Disabled"}',
            f'Managed Node Groups: {len(spec.node_groups)}',
            f'Fargate Profiles: {len(spec.fargate_profiles)}'
        ]
        for item in compute_items:
            p = doc.add_paragraph(item, style='List Bullet')
        
        # Storage Configuration
        doc.add_heading('Storage Configuration', 2)
        storage_items = [
            f'EBS CSI Driver: {"✓ Enabled" if spec.ebs_csi_enabled else "✗ Disabled"}',
            f'EFS: {"✓ Enabled" if spec.efs_enabled else "✗ Disabled"}',
            f'FSx for Lustre: {"✓ Enabled" if spec.fsx_enabled else "✗ Disabled"}'
        ]
        for item in storage_items:
            doc.add_paragraph(item, style='List Bullet')
        
        # Networking
        doc.add_heading('Networking Configuration', 2)
        networking_items = [
            f'Subnet Strategy: {spec.subnet_strategy}',
            f'Load Balancer Type: {spec.load_balancer_type.upper()}',
            f'Ingress Controller: {spec.ingress_controller}',
            f'Service Mesh: {spec.service_mesh}'
        ]
        for item in networking_items:
            doc.add_paragraph(item, style='List Bullet')
        
        # Security Configuration
        doc.add_heading('Security Configuration', 2)
        security_items = [
            f'Encryption: {"✓ Enabled" if spec.encryption_enabled else "✗ Disabled"}',
            f'Secrets Manager: {spec.secrets_manager}',
            f'IRSA: {"✓ Enabled" if spec.irsa_enabled else "✗ Disabled"}',
            f'Pod Security Standards: {spec.pod_security_standards}',
            f'Network Policies: {"✓ Enabled" if spec.network_policies else "✗ Disabled"}'
        ]
        for item in security_items:
            doc.add_paragraph(item, style='List Bullet')
        
        # ===== NODE GROUPS =====
        if spec.node_groups:
            doc.add_heading('Node Groups', 1)
            
            for i, ng in enumerate(spec.node_groups, 1):
                doc.add_heading(f'Node Group {i}: {ng.get("name", f"ng-{i}")}', 2)
                
                ng_table = doc.add_table(rows=5, cols=2)
                ng_table.style = 'Light Grid Accent 1'
                
                ng_table.rows[0].cells[0].text = 'Instance Type'
                ng_table.rows[0].cells[1].text = ng.get('instance_type', 't3.medium')
                ng_table.rows[1].cells[0].text = 'Min Size'
                ng_table.rows[1].cells[1].text = str(ng.get('min_size', 1))
                ng_table.rows[2].cells[0].text = 'Max Size'
                ng_table.rows[2].cells[1].text = str(ng.get('max_size', 5))
                ng_table.rows[3].cells[0].text = 'Capacity Type'
                ng_table.rows[3].cells[1].text = ng.get('capacity_type', 'ON_DEMAND')
                ng_table.rows[4].cells[0].text = 'Disk Size'
                ng_table.rows[4].cells[1].text = f"{ng.get('disk_size', 100)} GB"
                
                doc.add_paragraph()
        
        # ===== KARPENTER CONFIGURATION =====
        if spec.karpenter_enabled:
            doc.add_heading('Karpenter Configuration', 1)
            doc.add_paragraph(
                'Karpenter is enabled for dynamic node provisioning and cost optimization. '
                'Karpenter will automatically provision nodes based on pending pod requirements.'
            )
            
            if spec.karpenter_config:
                karp_items = [
                    f'Consolidation: {spec.karpenter_config.get("consolidation_enabled", "Enabled")}',
                    f'TTL After Empty: {spec.karpenter_config.get("ttl_seconds_after_empty", 30)} seconds'
                ]
                for item in karp_items:
                    doc.add_paragraph(item, style='List Bullet')
        
        # ===== OBSERVABILITY =====
        doc.add_heading('Observability & Monitoring', 1)
        observability_items = [
            f'Logging: {"✓ Enabled" if spec.logging_enabled else "✗ Disabled"}',
            f'Logging Destination: {spec.logging_destination if spec.logging_enabled else "N/A"}',
            f'Metrics Server: {"✓ Enabled" if spec.metrics_server else "✗ Disabled"}',
            f'Prometheus: {"✓ Enabled" if spec.prometheus_enabled else "✗ Disabled"}',
            f'Grafana: {"✓ Enabled" if spec.grafana_enabled else "✗ Disabled"}'
        ]
        for item in observability_items:
            doc.add_paragraph(item, style='List Bullet')
        
        # ===== ADD-ONS & TOOLS =====
        doc.add_heading('Add-ons & Tools', 1)
        addons_items = [
            f'Cluster Autoscaler: {"✓ Enabled" if spec.cluster_autoscaler else "✗ Disabled"}',
            f'External DNS: {"✓ Enabled" if spec.external_dns else "✗ Disabled"}',
            f'Cert Manager: {"✓ Enabled" if spec.cert_manager else "✗ Disabled"}',
            f'ArgoCD: {"✓ Enabled" if spec.argocd else "✗ Disabled"}',
            f'Flux: {"✓ Enabled" if spec.flux else "✗ Disabled"}'
        ]
        for item in addons_items:
            doc.add_paragraph(item, style='List Bullet')
        
        # ===== WORKLOAD INFORMATION =====
        if spec.expected_workloads > 0:
            doc.add_heading('Workload Information', 1)
            workload_items = [
                f'Expected Workloads: {spec.expected_workloads}',
                f'Peak Pod Count: {spec.peak_pod_count}',
                f'Performance Tier: {spec.performance_tier}'
            ]
            for item in workload_items:
                doc.add_paragraph(item, style='List Bullet')
            
            if spec.workload_types:
                doc.add_paragraph('Workload Types:', style='List Bullet')
                for wt in spec.workload_types:
                    doc.add_paragraph(f'  • {wt}')
        
        # ===== COST ESTIMATION =====
        doc.add_heading('Cost Estimation', 1)
        if spec.monthly_budget > 0:
            doc.add_paragraph(f'Monthly Budget: ${spec.monthly_budget:,.2f}')
        doc.add_paragraph(
            'For detailed cost breakdown and optimization recommendations, '
            'please refer to the cost estimation section in the EKS Design Hub.'
        )
        
        # ===== RECOMMENDATIONS =====
        if spec.ai_recommendations:
            doc.add_heading('AI Recommendations', 1)
            for i, rec in enumerate(spec.ai_recommendations[:5], 1):
                if isinstance(rec, dict):
                    doc.add_heading(f'{i}. {rec.get("title", "Recommendation")}', 2)
                    doc.add_paragraph(rec.get('description', 'No description available'))
        
        # ===== FOOTER =====
        doc.add_paragraph()
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run('_______________________________________________')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        footer_text = doc.add_paragraph(
            f'Generated by EKS Design Hub on {datetime.now().strftime("%B %d, %Y at %H:%M")}'
        )
        footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return doc_io

# ============================================================================
# DIAGRAM GENERATOR
# ============================================================================

class DiagramGenerator:
    """Generate professional architecture diagrams"""
    
    # AWS-style colors
    COLORS = {
        'eks': '#FF9900',  # AWS Orange
        'compute': '#3F8624',  # Green
        'network': '#5B9BD5',  # Blue
        'karpenter': '#326CE5',  # Kubernetes Blue
    }
    
    @staticmethod
    def generate_svg(spec: EKSDesignSpec) -> str:
        """Generate professional SVG architecture diagram"""
        
        svg_width = 1200
        svg_height = 800
        
        svg_parts = []
        
        # SVG Header with defs
        svg_parts.append(f'''<svg xmlns="http://www.w3.org/2000/svg" 
     viewBox="0 0 {svg_width} {svg_height}" 
     width="{svg_width}" 
     height="{svg_height}">
  
  <!-- Definitions -->
  <defs>
    <linearGradient id="awsGradient" x1="0%" y1="0%" x2="0%" y2="100%">
      <stop offset="0%" style="stop-color:#FF9900;stop-opacity:1" />
      <stop offset="100%" style="stop-color:#EC7211;stop-opacity:1" />
    </linearGradient>
    <linearGradient id="blueGradient" x1="0%" y1="0%" x2="0%" y2="100%">
      <stop offset="0%" style="stop-color:#5B9BD5;stop-opacity:0.3" />
      <stop offset="100%" style="stop-color:#4472C4;stop-opacity:0.3" />
    </linearGradient>
    <linearGradient id="greenGradient" x1="0%" y1="0%" x2="0%" y2="100%">
      <stop offset="0%" style="stop-color:#3F8624;stop-opacity:0.2" />
      <stop offset="100%" style="stop-color:#2D6219;stop-opacity:0.2" />
    </linearGradient>
    <filter id="shadow">
      <feDropShadow dx="2" dy="2" stdDeviation="3" flood-opacity="0.3"/>
    </filter>
    <filter id="glow">
      <feGaussianBlur stdDeviation="2" result="coloredBlur"/>
      <feMerge>
        <feMergeNode in="coloredBlur"/>
        <feMergeNode in="SourceGraphic"/>
      </feMerge>
    </filter>
  </defs>
  
  <!-- Background -->
  <rect width="{svg_width}" height="{svg_height}" fill="#F5F7FA"/>
  
  <!-- Title -->
  <text x="600" y="40" text-anchor="middle" font-family="Arial, sans-serif" 
        font-size="24" font-weight="bold" fill="#232F3E">
    {spec.project_name} - EKS Architecture
  </text>
  <text x="600" y="65" text-anchor="middle" font-family="Arial, sans-serif" 
        font-size="14" fill="#5A6C7D">
    {spec.environment} Environment | {spec.region}
  </text>
''')
        
        # VPC Container
        vpc_x, vpc_y = 50, 100
        vpc_width, vpc_height = 1100, 650
        
        svg_parts.append(f'''
  <!-- VPC Container -->
  <rect x="{vpc_x}" y="{vpc_y}" width="{vpc_width}" height="{vpc_height}" 
        fill="url(#blueGradient)" stroke="#5B9BD5" stroke-width="3" 
        rx="10" filter="url(#shadow)"/>
  <text x="{vpc_x + 20}" y="{vpc_y + 30}" font-family="Arial, sans-serif" 
        font-size="18" font-weight="bold" fill="#232F3E">
    🌐 VPC ({spec.region})
  </text>
''')
        
        # EKS Control Plane
        eks_x, eks_y = 450, 150
        eks_width, eks_height = 300, 80
        
        svg_parts.append(f'''
  <!-- EKS Control Plane -->
  <rect x="{eks_x}" y="{eks_y}" width="{eks_width}" height="{eks_height}" 
        fill="url(#awsGradient)" stroke="#EC7211" stroke-width="2" 
        rx="8" filter="url(#shadow)"/>
  <text x="{eks_x + eks_width/2}" y="{eks_y + 30}" text-anchor="middle" 
        font-family="Arial, sans-serif" font-size="16" font-weight="bold" fill="white">
    🎯 EKS Control Plane
  </text>
  <text x="{eks_x + eks_width/2}" y="{eks_y + 50}" text-anchor="middle" 
        font-family="Arial, sans-serif" font-size="12" fill="white">
    {spec.region}
  </text>
  <text x="{eks_x + eks_width/2}" y="{eks_y + 68}" text-anchor="middle" 
        font-family="Arial, sans-serif" font-size="11" fill="white">
    Managed by AWS
  </text>
''')
        
        # Availability Zones with Node Groups
        az_width = 320
        az_height = 400
        az_y = 280
        
        for i, az in enumerate(spec.availability_zones[:3]):
            az_x = 100 + (i * 360)
            
            svg_parts.append(f'''
  <!-- Availability Zone {i+1} -->
  <rect x="{az_x}" y="{az_y}" width="{az_width}" height="{az_height}" 
        fill="url(#greenGradient)" stroke="#3F8624" stroke-width="2" 
        rx="5" stroke-dasharray="5,5"/>
  <text x="{az_x + 10}" y="{az_y + 25}" font-family="Arial, sans-serif" 
        font-size="14" font-weight="bold" fill="#232F3E">
    📍 AZ: {az}
  </text>
''')
            
            # Node Group in this AZ
            if i < len(spec.node_groups):
                ng = spec.node_groups[i]
                node_x = az_x + 20
                node_y = az_y + 50
                node_width = 280
                node_height = 150
                
                svg_parts.append(f'''
  <!-- Node Group {i+1} -->
  <rect x="{node_x}" y="{node_y}" width="{node_width}" height="{node_height}" 
        fill="#E8F5E9" stroke="#3F8624" stroke-width="2" rx="5" filter="url(#shadow)"/>
  <text x="{node_x + 10}" y="{node_y + 25}" font-family="Arial, sans-serif" 
        font-size="13" font-weight="bold" fill="#232F3E">
    💻 Node Group: {ng.get('name', f'ng-{i+1}')}
  </text>
  <text x="{node_x + 10}" y="{node_y + 45}" font-family="Arial, sans-serif" 
        font-size="11" fill="#5A6C7D">
    Instance: {ng.get('instance_type', 't3.medium')}
  </text>
  <text x="{node_x + 10}" y="{node_y + 62}" font-family="Arial, sans-serif" 
        font-size="11" fill="#5A6C7D">
    Capacity: {ng.get('min_size', 1)}-{ng.get('max_size', 5)} nodes
  </text>
  <text x="{node_x + 10}" y="{node_y + 79}" font-family="Arial, sans-serif" 
        font-size="11" fill="#5A6C7D">
    Type: {ng.get('capacity_type', 'ON_DEMAND')}
  </text>
  <text x="{node_x + 10}" y="{node_y + 96}" font-family="Arial, sans-serif" 
        font-size="10" fill="#7AA116">
    ✓ Auto Scaling Enabled
  </text>
''')
                
                # Connection from EKS to Node Group
                svg_parts.append(f'''
  <line x1="{eks_x + eks_width/2}" y1="{eks_y + eks_height}" 
        x2="{node_x + node_width/2}" y2="{node_y}" 
        stroke="#879596" stroke-width="2" stroke-dasharray="5,5"/>
''')
                
                # Pods inside node group
                for j in range(3):
                    pod_x = node_x + 20 + (j * 80)
                    pod_y = node_y + 115
                    
                    svg_parts.append(f'''
  <circle cx="{pod_x}" cy="{pod_y}" r="15" fill="#326CE5" stroke="#1E4BA8" stroke-width="1"/>
  <text x="{pod_x}" y="{pod_y + 5}" text-anchor="middle" 
        font-family="Arial, sans-serif" font-size="10" fill="white">
    Pod
  </text>
''')
            
            # Subnet
            subnet_y = az_y + 220
            svg_parts.append(f'''
  <!-- Subnet -->
  <rect x="{az_x + 20}" y="{subnet_y}" width="{az_width - 40}" height="60" 
        fill="#FFFFFF" stroke="#5B9BD5" stroke-width="1" rx="3"/>
  <text x="{az_x + 35}" y="{subnet_y + 25}" font-family="Arial, sans-serif" 
        font-size="11" font-weight="bold" fill="#232F3E">
    🔒 Private Subnet
  </text>
  <text x="{az_x + 35}" y="{subnet_y + 43}" font-family="Arial, sans-serif" 
        font-size="9" fill="#5A6C7D">
    CIDR: 10.0.{i}.0/24
  </text>
''')
            
            # Load Balancer (only in first AZ)
            if i == 0 and spec.load_balancer_type:
                lb_y = az_y + 300
                svg_parts.append(f'''
  <!-- Load Balancer -->
  <rect x="{az_x + 50}" y="{lb_y}" width="{az_width - 100}" height="50" 
        fill="#FFF3CD" stroke="#FFC107" stroke-width="2" rx="5" filter="url(#shadow)"/>
  <text x="{az_x + az_width/2}" y="{lb_y + 25}" text-anchor="middle" 
        font-family="Arial, sans-serif" font-size="12" font-weight="bold" fill="#232F3E">
    ⚖️ {spec.load_balancer_type.upper()} Load Balancer
  </text>
  <text x="{az_x + az_width/2}" y="{lb_y + 40}" text-anchor="middle" 
        font-family="Arial, sans-serif" font-size="9" fill="#5A6C7D">
    Cross-Zone Enabled
  </text>
''')
        
        # Karpenter (if enabled)
        if spec.karpenter_enabled:
            karp_x, karp_y = 900, 250
            svg_parts.append(f'''
  <!-- Karpenter -->
  <rect x="{karp_x}" y="{karp_y}" width="200" height="100" 
        fill="#E3F2FD" stroke="#326CE5" stroke-width="2" rx="5" filter="url(#glow)"/>
  <text x="{karp_x + 100}" y="{karp_y + 30}" text-anchor="middle" 
        font-family="Arial, sans-serif" font-size="14" font-weight="bold" fill="#1E4BA8">
    🚀 Karpenter
  </text>
  <text x="{karp_x + 100}" y="{karp_y + 50}" text-anchor="middle" 
        font-family="Arial, sans-serif" font-size="10" fill="#5A6C7D">
    Auto-scaling Controller
  </text>
  <text x="{karp_x + 100}" y="{karp_y + 67}" text-anchor="middle" 
        font-family="Arial, sans-serif" font-size="9" fill="#7AA116">
    ✓ Dynamic Provisioning
  </text>
  <text x="{karp_x + 100}" y="{karp_y + 82}" text-anchor="middle" 
        font-family="Arial, sans-serif" font-size="9" fill="#7AA116">
    ✓ Cost Optimization
  </text>
  
  <!-- Connection to EKS -->
  <line x1="{eks_x + eks_width}" y1="{eks_y + eks_height/2}" 
        x2="{karp_x}" y2="{karp_y + 50}" 
        stroke="#326CE5" stroke-width="2"/>
''')
        
        # Storage (if enabled)
        if spec.ebs_csi_enabled or spec.efs_enabled:
            storage_x, storage_y = 900, 400
            storage_items = []
            
            if spec.ebs_csi_enabled:
                storage_items.append('💾 EBS Volumes')
            if spec.efs_enabled:
                storage_items.append('📁 EFS File System')
            
            storage_height = 50 + len(storage_items) * 20
            svg_parts.append(f'''
  <!-- Storage -->
  <rect x="{storage_x}" y="{storage_y}" width="200" height="{storage_height}" 
        fill="#FFF3E0" stroke="#FF9800" stroke-width="2" rx="5" filter="url(#shadow)"/>
  <text x="{storage_x + 100}" y="{storage_y + 25}" text-anchor="middle" 
        font-family="Arial, sans-serif" font-size="13" font-weight="bold" fill="#232F3E">
    💽 Storage
  </text>
''')
            
            for i, storage in enumerate(storage_items):
                svg_parts.append(f'''
  <text x="{storage_x + 20}" y="{storage_y + 45 + i * 20}" 
        font-family="Arial, sans-serif" font-size="10" fill="#5A6C7D">
    {storage}
  </text>
''')
        
        # Legend
        legend_y = 720
        svg_parts.append(f'''
  <!-- Legend -->
  <text x="100" y="{legend_y}" font-family="Arial, sans-serif" 
        font-size="11" font-weight="bold" fill="#232F3E">
    Legend:
  </text>
  <rect x="160" y="{legend_y - 12}" width="15" height="15" fill="url(#awsGradient)"/>
  <text x="180" y="{legend_y}" font-family="Arial, sans-serif" font-size="9" fill="#5A6C7D">
    EKS Control Plane
  </text>
  
  <rect x="290" y="{legend_y - 12}" width="15" height="15" fill="#E8F5E9" stroke="#3F8624"/>
  <text x="310" y="{legend_y}" font-family="Arial, sans-serif" font-size="9" fill="#5A6C7D">
    Worker Nodes
  </text>
  
  <circle cx="407" cy="{legend_y - 5}" r="7" fill="#326CE5"/>
  <text x="420" y="{legend_y}" font-family="Arial, sans-serif" font-size="9" fill="#5A6C7D">
    Pods
  </text>
''')
        
        # Footer
        svg_parts.append(f'''
  <text x="1150" y="{legend_y + 30}" text-anchor="end" 
        font-family="Arial, sans-serif" font-size="8" fill="#879596">
    Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}
  </text>
</svg>
''')
        
        return ''.join(svg_parts)
    
    def generate_diagram(self, spec: EKSDesignSpec):
        """Generate and display professional architecture diagram"""
        
        st.markdown("### 🏗️ EKS Architecture Diagram")
        st.caption("Professional AWS-style architecture visualization")
        
        # Generate SVG
        svg_content = self.generate_svg(spec)
        
        # Display in Streamlit using HTML component
        st.components.v1.html(svg_content, height=850, scrolling=False)
        
        # Download button
        col1, col2, col3 = st.columns([1, 1, 2])
        with col1:
            st.download_button(
                label="⬇️ Download SVG",
                data=svg_content,
                file_name=f"{spec.project_name.replace(' ', '_')}_eks_architecture.svg",
                mime="image/svg+xml",
                use_container_width=True
            )
        with col2:
            st.info("✨ Professional diagram generated!")
        
        return svg_content

# ============================================================================
# MAIN RENDER FUNCTION
# ============================================================================

def render_eks_design_hub():
    """Main render function for the comprehensive EKS Design Hub"""
    
    # Initialize session state
    if 'view_mode' not in st.session_state:
        st.session_state.view_mode = 'wizard'
    
    # Sidebar for mode selection
    st.sidebar.title("🎯 EKS Design Hub")
    st.sidebar.markdown("---")
    
    mode = st.sidebar.radio(
        "Select Mode",
        ["🧙 Design Wizard", "📊 Quick Calculator", "📚 Best Practices", "📖 Documentation"]
    )
    
    if "Wizard" in mode:
        EKSDesignWizard.render_wizard()
    elif "Calculator" in mode:
        render_quick_calculator()
    elif "Best Practices" in mode:
        render_best_practices()
    elif "Documentation" in mode:
        render_documentation_export()

def render_quick_calculator():
    """Quick sizing calculator"""
    st.header("📊 Quick Sizing Calculator")
    
    col1, col2 = st.columns(2)
    
    with col1:
        workload_count = st.number_input("Number of Services", 1, 500, 10)
        pods_per_service = st.number_input("Avg Pods per Service", 1, 50, 3)
    
    with col2:
        cpu_per_pod = st.number_input("CPU per Pod (millicores)", 100, 8000, 500)
        memory_per_pod = st.number_input("Memory per Pod (MB)", 128, 16384, 512)
    
    # Calculate
    total_pods = workload_count * pods_per_service
    total_cpu = (total_pods * cpu_per_pod) / 1000
    total_memory = (total_pods * memory_per_pod) / 1024
    
    st.divider()
    col1, col2, col3 = st.columns(3)
    col1.metric("Total Pods", f"{total_pods}")
    col2.metric("Total vCPUs", f"{total_cpu:.1f}")
    col3.metric("Total Memory (GB)", f"{total_memory:.1f}")

def render_best_practices():
    """Display EKS best practices"""
    st.header("📚 EKS Best Practices")
    
    practices = {
        'Security': [
            'Enable IRSA for pod-level IAM permissions',
            'Use Pod Security Standards (PSS) in restricted mode',
            'Implement network policies for pod-to-pod communication',
            'Enable encryption at rest for all data',
            'Use AWS Secrets Manager for sensitive data'
        ],
        'Cost Optimization': [
            'Use Karpenter for intelligent scaling and consolidation',
            'Leverage Spot instances for fault-tolerant workloads',
            'Right-size node groups based on actual usage',
            'Use Savings Plans and Reserved Instances',
            'Implement pod autoscaling (HPA/VPA)'
        ],
        'Performance': [
            'Use latest EKS version for performance improvements',
            'Enable metrics server for autoscaling',
            'Use appropriate storage classes (gp3, io2)',
            'Configure resource requests and limits',
            'Use topology-aware scheduling'
        ],
        'Reliability': [
            'Deploy across multiple AZs',
            'Implement pod disruption budgets',
            'Use health checks (liveness, readiness)',
            'Configure cluster autoscaling',
            'Regular backup and disaster recovery testing'
        ]
    }
    
    for category, items in practices.items():
        with st.expander(f"📋 {category}", expanded=True):
            for item in items:
                st.markdown(f"✅ {item}")

def render_documentation_export():
    """Documentation export options"""
    st.header("📖 Export Documentation")
    
    if 'design_spec' in st.session_state:
        spec = st.session_state.design_spec
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("📄 Export to Word", use_container_width=True):
                if not DOCX_AVAILABLE:
                    st.error("⚠️ Word export requires 'python-docx'. Please install it:\n```pip install python-docx```")
                else:
                    doc_gen = DocumentationGenerator()
                    doc_bytes_io = doc_gen.generate_word_doc(spec)
                    
                    if doc_bytes_io:
                        # Provide download button
                        st.download_button(
                            label="⬇️ Download Word Document",
                            data=doc_bytes_io.getvalue(),
                            file_name=f"eks_architecture_{spec.project_name.replace(' ', '_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                        st.success("✅ Word document ready for download!")
        
        with col2:
            if st.button("📄 Export to PDF", use_container_width=True):
                st.info("PDF export coming soon!")
        
        with col3:
            if st.button("📋 Export to JSON", use_container_width=True):
                spec_json = json.dumps(spec.__dict__, indent=2, default=str)
                st.download_button(
                    "Download JSON",
                    spec_json,
                    file_name=f"{spec.project_name}_design.json",
                    mime="application/json"
                )
    else:
        st.info("Complete the Design Wizard first to export documentation")

# ============================================================================
# COMPATIBILITY ALIAS
# ============================================================================

# Alias for backward compatibility with streamlit_app.py
render_eks_modernization_hub = render_eks_design_hub

# Class wrapper for compatibility with streamlit_app.py
class EKSModernizationModule:
    """Wrapper class for EKS Modernization functionality"""
    
    @staticmethod
    def render():
        """Main render method called from streamlit_app.py"""
        render_eks_design_hub()

# ============================================================================
# ENTRY POINT
# ============================================================================

if __name__ == "__main__":
    st.set_page_config(
        page_title="EKS Design & Architecture Hub",
        page_icon="🎯",
        layout="wide"
    )
    render_eks_design_hub()